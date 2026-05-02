#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any


def compact_text(value: Any, limit: int = 320) -> str:
    text = " ".join(str(value or "").split())
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "..."


def parse_docx(path: Path) -> dict[str, Any]:
    from docx import Document

    document = Document(str(path))
    title = path.stem
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Title") and title == path.stem:
            title = text
            continue

        if style_name.startswith("Heading"):
            if current:
                sections.append(finalize_docx_section(current))
            current = {
                "title": text,
                "style": style_name,
                "paragraphs": [],
            }
            continue

        if current is None:
            current = {
                "title": "Overview",
                "style": "Body",
                "paragraphs": [],
            }
            if title == path.stem:
                title = compact_text(text, 90)
        current["paragraphs"].append(text)

    if current:
        sections.append(finalize_docx_section(current))

    table_sections = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows[:8]:
            rows.append([compact_text(cell.text, 120) for cell in row.cells])
        table_sections.append(
            {
                "title": f"Table {table_index}",
                "dataType": "table",
                "summary": f"DOCX table with {len(table.rows)} row(s) and {len(table.columns)} column(s).",
                "columns": rows[0] if rows else [],
                "sampleRows": rows[1:],
                "rowCount": len(table.rows),
                "columnCount": len(table.columns),
            }
        )

    all_sections = sections + table_sections
    return {
        "documentType": "docx",
        "title": title,
        "sectionCount": len(all_sections),
        "paragraphCount": len(document.paragraphs),
        "tableCount": len(document.tables),
        "sections": all_sections
        or [
            {
                "title": "Overview",
                "dataType": "text",
                "summary": "No extractable DOCX content was found.",
                "bullets": [],
            }
        ],
    }


def finalize_docx_section(section: dict[str, Any]) -> dict[str, Any]:
    paragraphs = [p for p in section.get("paragraphs", []) if p]
    return {
        "title": compact_text(section.get("title"), 90),
        "dataType": "text",
        "style": section.get("style"),
        "summary": compact_text(" ".join(paragraphs[:2])),
        "bullets": [compact_text(p, 180) for p in paragraphs[:6]],
        "paragraphCount": len(paragraphs),
    }


def parse_xlsx(path: Path) -> dict[str, Any]:
    from openpyxl import load_workbook

    formulas_wb = load_workbook(str(path), data_only=False, read_only=False)
    values_wb = load_workbook(str(path), data_only=True, read_only=True)

    sections: list[dict[str, Any]] = []
    for sheet_name in formulas_wb.sheetnames:
        formula_sheet = formulas_wb[sheet_name]
        value_sheet = values_wb[sheet_name]
        rows = collect_non_empty_rows(value_sheet, max_rows=12)
        formula_cells = collect_formula_cells(formula_sheet, max_cells=40)
        tables = list(getattr(formula_sheet, "tables", {}).keys())
        chart_count = len(getattr(formula_sheet, "_charts", []) or [])

        if not rows and not formula_cells:
            continue

        headers = infer_headers(rows)
        sample_rows = rows[1:8] if len(rows) > 1 else []
        numeric_columns = infer_numeric_columns(headers, sample_rows)
        data_type = "chart" if numeric_columns else "table"

        sections.append(
            {
                "title": sheet_name,
                "dataType": data_type,
                "summary": summarize_sheet(headers, sample_rows, numeric_columns, formula_cells, tables, chart_count),
                "columns": headers,
                "sampleRows": sample_rows,
                "numericColumns": numeric_columns,
                "formulaCells": formula_cells,
                "tables": tables,
                "chartCount": chart_count,
                "maxRow": formula_sheet.max_row,
                "maxColumn": formula_sheet.max_column,
            }
        )

    props = formulas_wb.properties
    return {
        "documentType": "xlsx",
        "title": props.title or path.stem,
        "sheetCount": len(formulas_wb.sheetnames),
        "sheets": formulas_wb.sheetnames,
        "sections": sections
        or [
            {
                "title": "Workbook Overview",
                "dataType": "table",
                "summary": "No usable workbook rows or formulas were found.",
                "columns": [],
                "sampleRows": [],
                "numericColumns": [],
            }
        ],
    }


def collect_non_empty_rows(sheet: Any, max_rows: int) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for raw_row in sheet.iter_rows(values_only=True):
        normalized = [normalize_cell(cell) for cell in raw_row]
        if any(cell not in (None, "") for cell in normalized):
            rows.append(normalized)
        if len(rows) >= max_rows:
            break
    return rows


def collect_formula_cells(sheet: Any, max_cells: int) -> list[dict[str, str]]:
    cells: list[dict[str, str]] = []
    for row in sheet.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith("="):
                cells.append({"cell": cell.coordinate, "formula": value})
                if len(cells) >= max_cells:
                    return cells
    return cells


def normalize_cell(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (int, float, bool)):
        return value
    return compact_text(value, 160)


def infer_headers(rows: list[list[Any]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows[:3])
    first = rows[0]
    headers = []
    for index in range(width):
        value = first[index] if index < len(first) else None
        headers.append(str(value).strip() if value not in (None, "") else f"Column {index + 1}")
    return headers


def infer_numeric_columns(headers: list[str], sample_rows: list[list[Any]]) -> list[str]:
    numeric_columns = []
    for index, header in enumerate(headers):
        values = [row[index] for row in sample_rows if index < len(row)]
        numeric_count = sum(1 for value in values if isinstance(value, (int, float)) and not isinstance(value, bool))
        if numeric_count >= 2:
            numeric_columns.append(header)
    return numeric_columns


def summarize_sheet(
    headers: list[str],
    sample_rows: list[list[Any]],
    numeric_columns: list[str],
    formula_cells: list[dict[str, str]],
    tables: list[str],
    chart_count: int,
) -> str:
    parts = [f"{len(sample_rows)} sampled data row(s)"]
    if headers:
        parts.append(f"headers: {', '.join(headers[:5])}")
    if numeric_columns:
        parts.append(f"numeric columns: {', '.join(numeric_columns[:5])}")
    if formula_cells:
        parts.append(f"{len(formula_cells)} sampled formula cell(s)")
    if tables:
        parts.append(f"tables: {', '.join(tables[:3])}")
    if chart_count:
        parts.append(f"{chart_count} embedded chart(s)")
    return compact_text("; ".join(parts), 420)


def main() -> int:
    parser = argparse.ArgumentParser(description="Profile DOCX or XLSX source content for a Prompt-to-PPT pipeline.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", type=Path, help="Write JSON output to this path.")
    args = parser.parse_args()

    path = args.input
    if not path.exists():
        raise SystemExit(f"Input not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".docx":
        result = parse_docx(path)
    elif suffix == ".xlsx":
        result = parse_xlsx(path)
    else:
        raise SystemExit(f"Unsupported source type: {suffix}")

    payload = json.dumps(result, ensure_ascii=False, indent=2, default=str)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload, encoding="utf-8")
    else:
        print(payload)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
